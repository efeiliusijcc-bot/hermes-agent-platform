from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

MAX_EXTRACTED_CHARS = 2_000_000
MAX_PDF_PAGES = 500
MAX_EXCEL_CELLS = 100_000
MAX_ARCHIVE_ENTRIES = 10_000
MAX_ARCHIVE_UNCOMPRESSED_BYTES = 50_000_000
SUPPORTED_EXTENSIONS = {".md", ".markdown", ".txt", ".pdf", ".docx", ".xlsx"}


class DocumentParseError(ValueError):
    def __init__(self, message: str, *, status_code: int = 422) -> None:
        super().__init__(message)
        self.status_code = status_code


@dataclass(frozen=True)
class ParsedDocument:
    parser: str
    text: str


def parse_document(filename: str, content: bytes) -> ParsedDocument:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentParseError(
            "Supported document types are Markdown, text, PDF, Word .docx, and Excel .xlsx",
            status_code=415,
        )
    try:
        if extension in {".md", ".markdown", ".txt"}:
            parsed = ParsedDocument(parser="utf8-text", text=content.decode("utf-8"))
        elif extension == ".pdf":
            parsed = ParsedDocument(parser="pypdf", text=_parse_pdf(content))
        elif extension == ".docx":
            _validate_office_archive(content)
            parsed = ParsedDocument(parser="python-docx", text=_parse_docx(content))
        else:
            _validate_office_archive(content)
            parsed = ParsedDocument(parser="openpyxl", text=_parse_xlsx(content))
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"Document could not be parsed as {extension}") from exc

    text = parsed.text.replace("\x00", "").strip()
    if not text:
        raise DocumentParseError("Document contains no extractable text")
    if len(text) > MAX_EXTRACTED_CHARS:
        raise DocumentParseError("Extracted document text exceeds the configured limit", status_code=413)
    return ParsedDocument(parser=parsed.parser, text=text)


def chunk_text(text: str, *, max_chars: int, overlap_chars: int) -> list[str]:
    if max_chars < 256 or overlap_chars < 0 or overlap_chars >= max_chars:
        raise ValueError("invalid chunk configuration")
    normalized = "\n".join(line.rstrip() for line in text.replace("\r\n", "\n").split("\n")).strip()
    chunks: list[str] = []
    start = 0
    while start < len(normalized):
        hard_end = min(len(normalized), start + max_chars)
        end = hard_end
        if hard_end < len(normalized):
            boundary = max(normalized.rfind("\n", start + max_chars // 2, hard_end), normalized.rfind(" ", start + max_chars // 2, hard_end))
            if boundary > start:
                end = boundary
        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(normalized):
            break
        start = max(start + 1, end - overlap_chars)
    return chunks


def _parse_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    if len(reader.pages) > MAX_PDF_PAGES:
        raise DocumentParseError("PDF page count exceeds the configured limit", status_code=413)
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _parse_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _parse_xlsx(content: bytes) -> str:
    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    parts: list[str] = []
    cells_seen = 0
    try:
        for worksheet in workbook.worksheets:
            parts.append(f"# Sheet: {worksheet.title}")
            for row in worksheet.iter_rows(values_only=True):
                cells_seen += len(row)
                if cells_seen > MAX_EXCEL_CELLS:
                    raise DocumentParseError("Excel cell count exceeds the configured limit", status_code=413)
                values = ["" if value is None else str(value) for value in row]
                if any(values):
                    parts.append("\t".join(values))
    finally:
        workbook.close()
    return "\n".join(parts)


def _validate_office_archive(content: bytes) -> None:
    try:
        with ZipFile(BytesIO(content)) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_ARCHIVE_ENTRIES:
                raise DocumentParseError("Office archive contains too many entries", status_code=413)
            if sum(entry.file_size for entry in entries) > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
                raise DocumentParseError("Office archive expands beyond the configured limit", status_code=413)
    except BadZipFile as exc:
        raise DocumentParseError("Office document is not a valid archive") from exc
